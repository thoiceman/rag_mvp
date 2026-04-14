from pathlib import Path
from langchain_core.documents import Document
import pymupdf4llm
from docx import Document as DocxDocument


class DocumentLoader:
    def load_file(self, file_path: str) -> list[Document]:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix in [".txt", ".md"]:
            text = path.read_text(encoding="utf-8")
            return [Document(page_content=text, metadata={"source": path.name})]

        if suffix == ".pdf":
            # 使用 pymupdf4llm 将 PDF 高质量提取为 Markdown (支持表格、多栏等版面)
            md_text = pymupdf4llm.to_markdown(str(path))
            return [Document(page_content=md_text, metadata={"source": path.name})]

        if suffix == ".docx":
            doc = DocxDocument(str(path))
            lines = []
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                # 尝试将 Word 的标题样式转换为 Markdown 标题
                style_name = p.style.name.lower()
                if style_name.startswith('heading'):
                    try:
                        level = int(style_name.replace('heading ', ''))
                        lines.append(f"{'#' * level} {text}")
                    except ValueError:
                        lines.append(f"# {text}")
                else:
                    lines.append(text)
            
            md_text = "\n\n".join(lines)
            return [Document(page_content=md_text, metadata={"source": path.name})]

        raise ValueError(f"暂不支持的文件类型: {suffix}")