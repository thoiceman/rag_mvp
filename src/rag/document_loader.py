from pathlib import Path
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument


class DocumentLoader:
    def load_file(self, file_path: str) -> list[Document]:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix in [".txt", ".md"]:
            text = path.read_text(encoding="utf-8")
            return [Document(page_content=text, metadata={"source": path.name})]

        if suffix == ".pdf":
            reader = PdfReader(str(path))
            docs = []
            for i, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(page_content=text, metadata={"source": path.name, "page": i}))
            return docs

        if suffix == ".docx":
            doc = DocxDocument(str(path))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            return [Document(page_content=text, metadata={"source": path.name})]

        raise ValueError(f"暂不支持的文件类型: {suffix}")